# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "reportlab",
# ]
# ///

import os
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

def get_report_content():
    return [
        {"type": "title", "text": "E-Commerce Management Database System"},
        {"type": "title", "text": "Comprehensive Mini Project Report"},
        {"type": "normal", "text": "Course: CSE 303 Lab - Database Management (Summer 2026)"},
        {"type": "normal", "text": "Institution: Independent University, Bangladesh (IUB)"},
        {"type": "heading", "text": "Prepared By:"},
        {"type": "normal", "text": "- Asiful Islam (ID: 2420197, Section: 5)"},
        {"type": "normal", "text": "- Sakib Mohammad Ome (ID: 2220867, Section: 4)"},
        {"type": "normal", "text": "- Shadman Samin Shahriar (ID: 2211120, Section: 2)"},
        {"type": "heading", "text": "1. Introduction and System Overview"},
        {"type": "normal", "text": "The objective of this mini-project was to design, implement, and deploy a fully functional E-Commerce Database System, translating theoretical database management concepts into a practical, real-world software architecture."},
        {"type": "normal", "text": "Instead of relying merely on standalone SQL scripts, we integrated our database with a robust Django 5.2 backend and a responsive HTML5/Bootstrap 5 frontend. The database engine driving the application is a cloud-hosted Neon PostgreSQL 18.x instance. By strictly adhering to the requirements set forth in the CSE 303 Lab syllabus, we successfully implemented a cohesive platform that seamlessly manages inventory, users, financial transactions, and complex analytical reporting."},
        {"type": "heading", "text": "2. Database Design & Architecture (Task 1)"},
        {"type": "normal", "text": "Our first major milestone was establishing a robust Entity-Relationship (ER) model. The design strictly adhered to relational database principles, utilizing Crow's Foot notation to represent relationships, cardinalities, and entity dependencies."},
        {"type": "normal", "text": "We identified and designed 10 core tables to capture the operational reality of an e-commerce platform. These included: Category, Supplier, Product (seeded with Product IDs 101-113), Customer, Warehouse, Order, Order Detail (Weak Entity), Payment (1:1 Relationship with Order), Warehouse Stock (Bridge Table to resolve Many-to-Many), and an Order Log (Audit table)."},
        {"type": "normal", "text": "During the conceptualization phase, we refined the initial ER diagram to resolve structural flaws. We merged duplicate Customer tables into a single canonical entity, eliminated invalid direct links between the Order Log and Warehouse entities, and implemented bridge tables to properly satisfy Third Normal Form (3NF) requirements."},
        {"type": "heading", "text": "3. SQL Implementation and Analytical Queries (Task 2)"},
        {"type": "normal", "text": "Following the ER design, we generated the DDL schema. We utilized PostgreSQL SERIAL types for auto-incrementing primary keys and enforced strict data integrity using ON DELETE CASCADE and ON DELETE SET NULL constraints."},
        {"type": "normal", "text": "We then developed 10 mandatory analytical queries to extract actionable business intelligence from the dataset. These included pulling Customer Profiles, Product Details Catalogs, Customer Name Searches (pattern matching), Order Summary Reports using JOINs, Above-Average Price Products using Subqueries, Category Inventory Aggregations, High-Value Categories filtering, Warehouse Logistics Overviews, Pending Orders Tracking, and Customer Lifetime Spend Summaries."},
        {"type": "heading", "text": "4. Advanced Database Triggers (Task 2)"},
        {"type": "normal", "text": "To ensure data integrity at the database engine level, we programmed three PL/pgSQL Trigger Functions. 1) Stock Validation Trigger: Fires BEFORE INSERT on store_orderdetail to check requested quantity against available stock and raise an exception if it exceeds inventory. 2) Automated Payment Calculation Trigger: Fires BEFORE INSERT OR UPDATE on store_payment to automatically compute Final_Amount based on Amount, Tax, and Discount. 3) Order Deletion Audit Log Trigger: Fires AFTER DELETE on store_order to capture the deleted record and save a log into store_orderlog."},
        {"type": "heading", "text": "5. Scenario-Based Database Recovery — RAID Level 4 (Task 3)"},
        {"type": "normal", "text": "Data loss prevention was addressed through a RAID 4 storage simulation. We designed an architecture featuring 6 Data Disks (D1 through D6) and 1 dedicated Parity Disk (P). We used 4-bit data blocks representing core tables (e.g., D1=1010, D2=1100). Parity (P) was calculated using a bitwise XOR operation across all data disks resulting in 0101. If disk D5 suffers a hardware failure, the system reconstructs the lost block by XORing the surviving data disks against the Parity disk, successfully outputting the original 0011 block."},
        {"type": "heading", "text": "6. Database Normalization (Task 4)"},
        {"type": "normal", "text": "To prevent data anomalies, our schema was mathematically normalized. We walked an unnormalized sample table through First Normal Form (1NF) by flattening repeating groups to atomic values, Second Normal Form (2NF) by removing partial dependencies from composite keys, and Third Normal Form (3NF) by eliminating transitive dependencies (extracting Customer details away from Order IDs)."},
        {"type": "heading", "text": "7. Hashing, Indexing & B+ Trees (Task 5)"},
        {"type": "normal", "text": "To optimize search times for the store_product catalog, we implemented an Order-3 B+ Tree indexing structure on the Product_ID key, processing an insertion sequence of 13 keys. We simulated traversing the internal node routing paths to perform a fast Search for Product_ID 108, and simulated a Node Deletion for Product_ID 105, demonstrating how the tree automatically borrows from sibling nodes to rebalance and prevent underflow."},
        {"type": "heading", "text": "8. Web Application Integration (Task 6)"},
        {"type": "normal", "text": "The theoretical database design was brought to life via a Django web application (E_Com_Website). The application features User Authentication, a Product CRUD Management Dashboard for administrators, an Analytical Query Runner to execute the 10 SQL queries in real-time, and a Triggers and Theory Test Bench providing interactive forms to safely execute and observe the database triggers."},
        {"type": "heading", "text": "9. Advanced Cloud Media Integration (Phase 9)"},
        {"type": "normal", "text": "To support product image uploads on ephemeral cloud hosts, we integrated the Cloudinary SDK. We engineered a Deferred Upload Protocol where images are only transmitted to the Cloudinary CDN after the rest of the database form passes validation. This prevents orphaned files in the cloud storage bucket and saves bandwidth."},
        {"type": "heading", "text": "10. Production Deployment Architecture (Phase 10)"},
        {"type": "normal", "text": "The application was hardened for live production hosting on the Railway cloud platform. We configured Gunicorn as the multi-worker WSGI Server, WhiteNoise for zero-overhead static asset management, and nixpacks.toml for automated environment builds. Dynamic database routing was achieved using dj-database-url, allowing smooth failover between local SQLite and production Neon PostgreSQL clusters."},
        {"type": "heading", "text": "Conclusion"},
        {"type": "normal", "text": "This Mini Project represents a comprehensive, end-to-end database engineering effort. By marrying strict relational theory, normalization, and PL/pgSQL triggers with modern web frameworks and cloud infrastructure, we successfully developed a secure, scalable, and highly performant E-Commerce Database System."}
    ]

def generate_docx(filepath):
    doc = Document()
    content = get_report_content()
    for item in content:
        if item["type"] == "title":
            heading = doc.add_heading(item["text"], level=0)
            heading.alignment = 1 # Center
        elif item["type"] == "heading":
            doc.add_heading(item["text"], level=1)
        elif item["type"] == "normal":
            p = doc.add_paragraph(item["text"])
            p.alignment = 3 # Justify
    doc.save(filepath)

def generate_pdf(filepath):
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=14,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=14
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10,
        alignment=TA_JUSTIFY
    )
    
    story = []
    content = get_report_content()
    
    for item in content:
        if item["type"] == "title":
            story.append(Paragraph(item["text"], title_style))
            story.append(Spacer(1, 12))
        elif item["type"] == "heading":
            story.append(Paragraph(item["text"], heading_style))
        elif item["type"] == "normal":
            story.append(Paragraph(item["text"], normal_style))
            
    doc.build(story)

if __name__ == "__main__":
    out_dir = os.path.join("static", "reports")
    os.makedirs(out_dir, exist_ok=True)
    
    docx_path = os.path.join(out_dir, "Comprehensive_Project_Report.docx")
    pdf_path = os.path.join(out_dir, "Comprehensive_Project_Report.pdf")
    
    generate_docx(docx_path)
    print(f"Generated {docx_path}")
    
    generate_pdf(pdf_path)
    print(f"Generated {pdf_path}")
